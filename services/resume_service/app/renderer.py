import os
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import LETTER
from reportlab.pdfgen import canvas

from app.models import RenderResumeRequest, RenderResumeResponse

OUTPUT_DIR = Path(os.getenv("OUTPUT_DIR", "/app/output"))
SHOW_PROJECT_NAMES_IN_RESUME = os.getenv("SHOW_PROJECT_NAMES_IN_RESUME", "false").lower() == "true"


def _slugify(value: str) -> str:
    value = value.lower()
    value = re.sub(r"[^a-z0-9]+", "-", value)
    return value.strip("-")


def _estimate_keyword_coverage(jd_required_skills: list[str], resume_skills: list[str]) -> float:
    """Calculate keyword coverage based on JD required skills present in resume."""
    if not jd_required_skills:
        return 0.0
    
    resume_skills_lower = {s.lower() for s in resume_skills}
    matched = sum(1 for skill in jd_required_skills if skill.lower() in resume_skills_lower)
    return round(matched / len(jd_required_skills) * 100, 2)


def _estimate_ats_score(
    keyword_coverage_pct: float,
    bullet_count: int,
    has_experience_section: bool,
    skills_organized: bool,
    has_summary: bool
) -> float:
    """Calculate honest internal ATS alignment score.
    
    Factors:
    - Keyword coverage (40%)
    - Experience bullets quality (30%)
    - Resume structure (20%)
    - Summary presence (10%)
    
    Penalties:
    - Missing experience section: -20
    - Too few bullets (<3): -15
    - Keyword stuffing (>90% coverage with <4 bullets): -25
    """
    score = 0.0
    
    # Keyword coverage component (max 40 points)
    score += keyword_coverage_pct * 0.4
    
    # Experience bullets component (max 30 points)
    if bullet_count >= 5:
        score += 30
    elif bullet_count >= 3:
        score += 20
    elif bullet_count >= 1:
        score += 10
    else:
        score += 0
    
    # Structure component (max 20 points)
    if has_experience_section:
        score += 10
    if skills_organized:
        score += 10
    
    # Summary component (max 10 points)
    if has_summary:
        score += 10
    
    # Penalties
    if not has_experience_section:
        score -= 20
        print("[score] penalty: missing_experience_section=-20")
    
    if bullet_count < 3:
        score -= 15
        print(f"[score] penalty: too_few_bullets={bullet_count}, penalty=-15")
    
    # Keyword stuffing detection
    if keyword_coverage_pct > 90 and bullet_count < 4:
        score -= 25
        print(f"[score] penalty: keyword_stuffing (coverage={keyword_coverage_pct}%, bullets={bullet_count}), penalty=-25")
    
    # Log components
    print(f"[score] keyword_coverage={keyword_coverage_pct}%")
    print(f"[score] bullet_quality_score={bullet_count} bullets")
    print(f"[score] structure_score=experience:{has_experience_section}, skills:{skills_organized}")
    
    final_score = round(max(min(score, 100), 0), 2)
    print(f"[score] final_internal_ats_alignment={final_score}")
    return final_score


def _organize_skills_by_category(skills: list[str]) -> dict[str, list[str]]:
    """Organize skills into categories for better ATS structure."""
    categories = {
        "Languages": [],
        "Backend": [],
        "Frontend": [],
        "Cloud/DevOps": [],
        "Databases": [],
        "Tools": [],
    }
    
    # Categorization rules
    language_keywords = ['java', 'javascript', 'typescript', 'python', 'c#', 'node.js', 'c++']
    backend_keywords = ['spring', 'spring boot', 'express', 'expressjs', 'asp.net', 'rest api', 'graphql', 'fastapi', 'flask', 'django']
    frontend_keywords = ['react', 'reactjs', 'angular', 'vue', 'html', 'css', 'tailwind']
    cloud_keywords = ['aws', 'azure', 'docker', 'kubernetes', 'jenkins', 'circleci', 'terraform', 'ci/cd']
    database_keywords = ['mysql', 'mongodb', 'postgresql', 'nosql', 'redis', 'sql']
    tool_keywords = ['git', 'github', 'jira', 'confluence', 'postman', 'swagger', 'splunk', 'selenium']
    
    for skill in skills:
        skill_lower = skill.lower()
        categorized = False
        
        if any(kw in skill_lower for kw in language_keywords):
            categories["Languages"].append(skill)
            categorized = True
        elif any(kw in skill_lower for kw in backend_keywords):
            categories["Backend"].append(skill)
            categorized = True
        elif any(kw in skill_lower for kw in frontend_keywords):
            categories["Frontend"].append(skill)
            categorized = True
        elif any(kw in skill_lower for kw in cloud_keywords):
            categories["Cloud/DevOps"].append(skill)
            categorized = True
        elif any(kw in skill_lower for kw in database_keywords):
            categories["Databases"].append(skill)
            categorized = True
        elif any(kw in skill_lower for kw in tool_keywords):
            categories["Tools"].append(skill)
            categorized = True
        
        # If not categorized, add to Tools as fallback
        if not categorized:
            categories["Tools"].append(skill)
    
    # Remove empty categories
    return {k: v for k, v in categories.items() if v}


def render_resume(request: RenderResumeRequest) -> RenderResumeResponse:
    cp = request.candidate_profile
    jp = request.job_posting
    tc = request.tailored_content

    slug = _slugify(f"{cp.full_name}-{jp.company}-{jp.title}")
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    docx_path = OUTPUT_DIR / f"{slug}.docx"
    pdf_path = OUTPUT_DIR / f"{slug}.pdf"

    contact_parts = [x for x in [cp.email, cp.phone, cp.location] if x]
    contact_line = " | ".join(contact_parts)
    
    # Organize skills by category
    skills_to_use = tc.reordered_skills or cp.master_skills
    skill_categories = _organize_skills_by_category(skills_to_use)
    
    # Get education, certifications, projects, and employer groups from metadata
    education_lines = request.metadata.get('education_lines', [])
    certification_lines = request.metadata.get('certification_lines', [])
    experience_projects = request.metadata.get('experience_projects', [])
    employer_groups = request.metadata.get('employer_groups', [])
    
    print(f"[renderer] section_order=header,summary,experience,skills,education,certifications")
    print(f"[renderer] education_sections={len(education_lines)}")
    print(f"[renderer] certification_sections={len(certification_lines)}")
    print(f"[renderer] experience_projects={len(experience_projects)}")
    print(f"[renderer] employer_groups={len(employer_groups)}")
    print(f"[renderer] grouped_experience={len(employer_groups) > 0}")
    print(f"[resume] show_project_names={SHOW_PROJECT_NAMES_IN_RESUME}")
    print(f"[renderer] max_bullets_per_employer=6")
    print(f"[renderer] compact_resume_format=true")
    print(f"[renderer] orphan_bullet_avoidance=true")

    # --- DOCX ---
    doc = Document()
    
    # Set default font size for compact formatting
    style = doc.styles['Normal']
    font = style.font
    font.size = Pt(10.5)
    
    # A. Header
    name_para = doc.add_heading(cp.full_name, level=1)
    name_para.alignment = 1  # Center
    name_para.paragraph_format.space_after = Pt(3)
    
    if contact_line:
        contact_para = doc.add_paragraph(contact_line)
        contact_para.alignment = 1  # Center
        contact_para.paragraph_format.space_after = Pt(6)

    # B. Professional Summary
    summary_heading = doc.add_heading("Professional Summary", level=2)
    summary_heading.paragraph_format.space_after = Pt(3)
    summary_para = doc.add_paragraph(tc.summary)
    summary_para.paragraph_format.space_after = Pt(6)

    # C. Professional Experience (BEFORE Technical Skills)
    exp_heading = doc.add_heading("Professional Experience", level=2)
    exp_heading.paragraph_format.space_after = Pt(3)
    
    if employer_groups:
        # Use grouped employer structure
        for idx, group in enumerate(employer_groups):
            # Employer header
            employer_para = doc.add_paragraph()
            employer_run = employer_para.add_run(f"{group['employer']}")
            employer_run.bold = True
            employer_para.paragraph_format.space_after = Pt(2)
            
            # Role and Duration line
            role_para = doc.add_paragraph()
            if group.get('role'):
                role_run = role_para.add_run(group['role'])
                role_run.italic = True
            if group.get('duration'):
                if group.get('role'):
                    role_para.add_run(' | ')
                role_para.add_run(group['duration'])
            if group.get('location'):
                if group.get('role') or group.get('duration'):
                    role_para.add_run(' | ')
                role_para.add_run(group['location'])
            role_para.paragraph_format.space_after = Pt(3)
            
            # Merge all project bullets under employer (hide project names)
            if not SHOW_PROJECT_NAMES_IN_RESUME:
                all_bullets = []
                for proj in group.get('projects', []):
                    if proj.get('name'):
                        print(f"[resume] hidden_project_name={proj['name']}")
                    responsibilities = proj.get('responsibilities', [])
                    for resp in responsibilities:
                        if len(resp) > 20 and not resp.lower().startswith('worked as a full-stack developer following'):
                            if resp not in all_bullets:  # Deduplicate
                                all_bullets.append(resp)
                
                # Render merged bullets (limit to 6 strongest)
                for bullet in all_bullets[:6]:
                    bullet_para = doc.add_paragraph(bullet, style="List Bullet")
                    bullet_para.paragraph_format.space_after = Pt(2)
                
                print(f"[resume] merged_project_bullets_under_employer={group['employer']} count={len(all_bullets[:6])}")
            else:
                # Show project names (original behavior)
                for proj in group.get('projects', []):
                    if proj.get('name'):
                        proj_para = doc.add_paragraph()
                        proj_run = proj_para.add_run(f"Project: {proj['name']}")
                        proj_run.bold = True
                        if proj.get('duration'):
                            proj_para.add_run(f" | {proj['duration']}")
                    
                    responsibilities = proj.get('responsibilities', [])
                    for resp in responsibilities[:6]:
                        if len(resp) > 20 and not resp.lower().startswith('worked as a full-stack developer following'):
                            bullet_para = doc.add_paragraph(resp, style="List Bullet")
                            bullet_para.paragraph_format.space_after = Pt(2)
            
            # Reduced spacing between employers (avoid orphan bullets)
            if idx < len(employer_groups) - 1:
                spacer = doc.add_paragraph()
                spacer.paragraph_format.space_after = Pt(4)
    
    elif experience_projects and any(proj.get('responsibilities') for proj in experience_projects):
        # Fallback: use structured project data if responsibilities exist
        for proj in experience_projects:
            if proj.get('company') or proj.get('client'):
                company_name = proj.get('client') or proj.get('company')
                role = proj.get('role', 'Developer')
                duration = proj.get('duration', '')
                
                company_para = doc.add_paragraph()
                company_run = company_para.add_run(f"{company_name}")
                company_run.bold = True
                company_para.paragraph_format.space_after = Pt(2)
                
                if role or duration:
                    role_para = doc.add_paragraph()
                    if role:
                        role_run = role_para.add_run(role)
                        role_run.italic = True
                    if duration:
                        if role:
                            role_para.add_run(" | ")
                        role_para.add_run(duration)
                    role_para.paragraph_format.space_after = Pt(3)
            
            responsibilities = proj.get('responsibilities', [])
            for resp in responsibilities[:6]:
                if len(resp) > 20:
                    bullet_para = doc.add_paragraph(resp, style="List Bullet")
                    bullet_para.paragraph_format.space_after = Pt(2)
    
    elif tc.selected_bullets:
        # Fallback: use selected bullets
        for bullet in tc.selected_bullets:
            bullet_para = doc.add_paragraph(bullet, style="List Bullet")
            bullet_para.paragraph_format.space_after = Pt(2)
    else:
        doc.add_paragraph("Experience details available upon request.")

    # D. Technical Skills
    skills_heading = doc.add_heading("Technical Skills", level=2)
    skills_heading.paragraph_format.space_after = Pt(3)
    for category, skills in skill_categories.items():
        skills_text = f"{category}: {', '.join(skills)}"
        para = doc.add_paragraph(skills_text)
        # Compact spacing
        para.paragraph_format.space_after = Pt(2)

    # E. Education (if available)
    if education_lines:
        edu_heading = doc.add_heading("Education", level=2)
        edu_heading.paragraph_format.space_after = Pt(3)
        for edu_line in education_lines:
            # Clean up education lines
            if 'intermediate schooling from' in edu_line.lower():
                # Format: "Intermediate — International Public School"
                school = edu_line.split('from')[-1].strip()
                edu_para = doc.add_paragraph(f"Intermediate — {school}")
                edu_para.paragraph_format.space_after = Pt(2)
            elif 'sslc from' in edu_line.lower():
                # Format: "SSLC — Surendranath Centenary School"
                school = edu_line.split('from')[-1].strip()
                edu_para = doc.add_paragraph(f"SSLC — {school}")
                edu_para.paragraph_format.space_after = Pt(2)
            else:
                edu_para = doc.add_paragraph(edu_line)
                edu_para.paragraph_format.space_after = Pt(2)

    # F. Certifications (if available)
    if certification_lines:
        cert_heading = doc.add_heading("Certifications", level=2)
        cert_heading.paragraph_format.space_after = Pt(3)
        for cert_line in certification_lines:
            cert_para = doc.add_paragraph(cert_line, style="List Bullet")
            cert_para.paragraph_format.space_after = Pt(2)

    doc.save(str(docx_path))

    # --- PDF ---
    c = canvas.Canvas(str(pdf_path), pagesize=LETTER)
    width, height = LETTER
    y = height - 50

    def draw_line(text: str, font: str = "Helvetica", size: int = 10, gap: int = 14, indent: int = 0):
        nonlocal y
        c.setFont(font, size)
        # Handle long lines by wrapping
        max_width = 95 - (indent // 10)
        if len(text) > max_width:
            words = text.split()
            line = ""
            for word in words:
                if len(line + word) < max_width:
                    line += word + " "
                else:
                    c.drawString(50 + indent, y, line.strip())
                    y -= gap
                    if y < 50:  # Page break
                        c.showPage()
                        y = height - 50
                        c.setFont(font, size)
                    line = word + " "
            if line:
                c.drawString(50 + indent, y, line.strip())
                y -= gap
        else:
            c.drawString(50 + indent, y, text)
            y -= gap
        
        if y < 50:  # Page break
            c.showPage()
            y = height - 50

    def draw_heading(text: str):
        nonlocal y
        y -= 4
        draw_line(text, font="Helvetica-Bold", size=12, gap=16)

    # A. Header
    draw_line(cp.full_name, font="Helvetica-Bold", size=15, gap=18)
    if contact_line:
        draw_line(contact_line, size=9, gap=14)

    # B. Professional Summary
    draw_heading("Professional Summary")
    draw_line(tc.summary, size=10, gap=14)

    # C. Professional Experience
    draw_heading("Professional Experience")
    
    if employer_groups:
        # Use grouped employer structure
        for idx, group in enumerate(employer_groups):
            # Employer name
            draw_line(group['employer'], font="Helvetica-Bold", size=10, gap=12)
            
            # Role, Duration, Location
            role_parts = []
            if group.get('role'):
                role_parts.append(group['role'])
            if group.get('duration'):
                role_parts.append(group['duration'])
            if group.get('location'):
                role_parts.append(group['location'])
            
            if role_parts:
                draw_line(' | '.join(role_parts), font="Helvetica-Oblique", size=9, gap=12)
            
            # Merge all project bullets under employer (hide project names)
            if not SHOW_PROJECT_NAMES_IN_RESUME:
                all_bullets = []
                for proj in group.get('projects', []):
                    responsibilities = proj.get('responsibilities', [])
                    for resp in responsibilities:
                        if len(resp) > 20 and not resp.lower().startswith('worked as a full-stack developer following'):
                            if resp not in all_bullets:  # Deduplicate
                                all_bullets.append(resp)
                
                # Render merged bullets (limit to 6 strongest)
                for bullet in all_bullets[:6]:
                    draw_line(f"• {bullet}", size=9, gap=11, indent=10)
            else:
                # Show project names (original behavior)
                for proj in group.get('projects', []):
                    if proj.get('name'):
                        proj_name = f"Project: {proj['name']}"
                        if proj.get('duration'):
                            proj_name += f" | {proj['duration']}"
                        draw_line(proj_name, font="Helvetica-Bold", size=9, gap=11)
                    
                    responsibilities = proj.get('responsibilities', [])
                    for resp in responsibilities[:6]:
                        if len(resp) > 20 and not resp.lower().startswith('worked as a full-stack developer following'):
                            draw_line(f"• {resp}", size=9, gap=11, indent=10)
            
            # Reduced spacing between employers
            if idx < len(employer_groups) - 1:
                y -= 6
    
    elif experience_projects and any(proj.get('responsibilities') for proj in experience_projects):
        for proj in experience_projects:
            if proj.get('company') or proj.get('client'):
                company_name = proj.get('client') or proj.get('company')
                draw_line(company_name, font="Helvetica-Bold", size=10, gap=12)
                
                role = proj.get('role', '')
                duration = proj.get('duration', '')
                if role or duration:
                    role_line = f"{role}" if role else ""
                    if duration:
                        role_line += f" | {duration}" if role else duration
                    draw_line(role_line, font="Helvetica-Oblique", size=9, gap=12)
            
            responsibilities = proj.get('responsibilities', [])
            for resp in responsibilities[:6]:
                if len(resp) > 20:
                    draw_line(f"• {resp}", size=9, gap=11, indent=10)
            
            y -= 6
    
    elif tc.selected_bullets:
        for bullet in tc.selected_bullets:
            draw_line(f"• {bullet}", size=9, gap=11, indent=10)
    else:
        draw_line("Experience details available upon request.", size=10)

    # D. Technical Skills
    draw_heading("Technical Skills")
    for category, skills in skill_categories.items():
        skills_text = f"{category}: {', '.join(skills)}"
        draw_line(skills_text, size=9, gap=11)

    # E. Education
    if education_lines:
        draw_heading("Education")
        for edu_line in education_lines:
            # Clean up education lines
            if 'intermediate schooling from' in edu_line.lower():
                school = edu_line.split('from')[-1].strip()
                draw_line(f"Intermediate — {school}", size=9, gap=11)
            elif 'sslc from' in edu_line.lower():
                school = edu_line.split('from')[-1].strip()
                draw_line(f"SSLC — {school}", size=9, gap=11)
            else:
                draw_line(edu_line, size=9, gap=11)

    # F. Certifications
    if certification_lines:
        draw_heading("Certifications")
        for cert_line in certification_lines:
            draw_line(f"• {cert_line}", size=9, gap=11, indent=10)

    c.save()

    # Calculate honest ATS score
    jd_required = jp.required_skills or []
    keyword_coverage_pct = _estimate_keyword_coverage(jd_required, skills_to_use)
    
    has_experience_section = len(tc.selected_bullets) > 0 or len(experience_projects) > 0
    skills_organized = len(skill_categories) > 1
    has_summary = len(tc.summary) > 20
    
    ats_score_internal = _estimate_ats_score(
        keyword_coverage_pct=keyword_coverage_pct,
        bullet_count=len(tc.selected_bullets),
        has_experience_section=has_experience_section,
        skills_organized=skills_organized,
        has_summary=has_summary
    )

    return RenderResumeResponse(
        docx_path=str(docx_path),
        pdf_path=str(pdf_path),
        keyword_coverage_pct=keyword_coverage_pct,
        ats_score_internal=ats_score_internal,
    )
